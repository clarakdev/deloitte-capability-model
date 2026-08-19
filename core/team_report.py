"""
Word document generation for project Team Capability Reports.
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DELOITTE_GREEN = "86BC25"
DARK_GREY = "333333"
LIGHT_GREY = "F2F2F2"
MID_GREY = "666666"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_text(
    cell,
    text,
    *,
    bold=False,
    color=DARK_GREY,
    size=9,
) -> None:
    cell.text = ""

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))

    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)

    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(DELOITTE_GREEN)


def _add_small_label(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_after = Pt(3)

    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.name = "Arial"
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string(DARK_GREY)

    value_run = paragraph.add_run(value)
    value_run.font.name = "Arial"
    value_run.font.size = Pt(9)
    value_run.font.color.rgb = RGBColor.from_string(MID_GREY)


def _clean_join(values: list[str]) -> str:
    values = [str(value).strip() for value in values if str(value).strip()]
    return ", ".join(values) if values else "None recorded"


def _add_team_overview_table(
    document: Document,
    entries: list[dict],
) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = [
        "Project Role",
        "Assigned Employee",
        "Title",
        "Level",
        "Match",
    ]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, DARK_GREY)
        _set_cell_text(
            cell,
            header,
            bold=True,
            color="FFFFFF",
            size=9,
        )

    for entry in entries:
        employee = entry["employee"]

        cells = table.add_row().cells

        values = [
            entry["role_title"],
            employee.get("name", ""),
            employee.get("title", ""),
            employee.get("role_level", ""),
            f"{round(entry['match_score'] * 100)}%",
        ]

        for index, value in enumerate(values):
            _set_cell_text(cells[index], value)

    document.add_paragraph()


def _add_team_metrics(
    document: Document,
    average_team_match: int,
    roles_with_gaps: int,
    total_roles: int,
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    metric_values = [
        (
            f"{average_team_match}%",
            "Average Team Match",
        ),
        (
            f"{roles_with_gaps} of {total_roles}",
            "Roles With Gaps",
        ),
    ]

    for index, (value, label) in enumerate(metric_values):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, LIGHT_GREY)

        cell.text = ""

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        value_run = paragraph.add_run(value)
        value_run.bold = True
        value_run.font.name = "Arial"
        value_run.font.size = Pt(18)
        value_run.font.color.rgb = RGBColor.from_string(DELOITTE_GREEN)

        label_paragraph = cell.add_paragraph()
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        label_run = label_paragraph.add_run(label)
        label_run.font.name = "Arial"
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor.from_string(MID_GREY)

    document.add_paragraph()


def _add_fit_summary(
    document: Document,
    entry: dict,
) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    values = [
        (
            f"{round(entry['match_score'] * 100)}%",
            "Overall Match",
        ),
        (
            f"{entry['avg_fit']}/5",
            "Avg Fit",
        ),
        (
            str(entry["covered_count"]),
            "Skills Covered",
        ),
        (
            str(entry["gap_count"]),
            "Gaps to Address",
        ),
    ]

    for index, (value, label) in enumerate(values):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, LIGHT_GREY)

        cell.text = ""

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        value_run = paragraph.add_run(value)
        value_run.bold = True
        value_run.font.name = "Arial"
        value_run.font.size = Pt(14)
        value_run.font.color.rgb = RGBColor.from_string(DELOITTE_GREEN)

        label_paragraph = cell.add_paragraph()
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        label_run = label_paragraph.add_run(label)
        label_run.font.name = "Arial"
        label_run.font.size = Pt(7)
        label_run.font.color.rgb = RGBColor.from_string(MID_GREY)

    document.add_paragraph()


def _add_capability_table(
    document: Document,
    fit_report: list[dict],
) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = [
        "Required Capability",
        "Weight",
        "Closest Employee Skill",
        "Fit",
    ]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, DARK_GREY)
        _set_cell_text(
            cell,
            header,
            bold=True,
            color="FFFFFF",
            size=8,
        )

    for fit in fit_report:
        cells = table.add_row().cells

        similarity = float(fit.get("similarity", 0.0))
        fit_score = min(5, max(1, int(similarity * 5 + 0.999999)))

        values = [
            fit.get("cap_name", ""),
            fit.get("weight", ""),
            fit.get("best_match_skill") or "No match found",
            f"{fit_score}/5",
        ]

        for index, value in enumerate(values):
            color = (
                "C00000"
                if fit.get("is_gap") and index == 3
                else DARK_GREY
            )

            _set_cell_text(
                cells[index],
                value,
                color=color,
                size=8,
            )

    document.add_paragraph()


def build_team_report_docx(
    project: dict,
    entries: list[dict],
    team_summary: str,
) -> BytesIO:
    """
    Build the final Team Capability Report and return it as an in-memory DOCX.
    """
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    # ── Cover / project overview ──────────────────────────────────────────

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    title_run = title.add_run(project.get("name", "Project"))
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor.from_string(DARK_GREY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run("TEAM CAPABILITY REPORT")
    subtitle_run.bold = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string(DELOITTE_GREEN)

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_run = date_paragraph.add_run(
        f"Generated {datetime.now().strftime('%d %B %Y')}"
    )
    date_run.font.name = "Arial"
    date_run.font.size = Pt(8)
    date_run.font.color.rgb = RGBColor.from_string(MID_GREY)

    _add_section_heading(document, "Project Overview")

    description = document.add_paragraph(
        project.get("description", "") or "No project description provided."
    )
    description.paragraph_format.space_after = Pt(10)
    description.paragraph_format.line_spacing = 1.15

    for run in description.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(DARK_GREY)

    if project.get("client"):
        _add_small_label(
            document,
            "Client",
            str(project["client"]),
        )

    # ── Team overview ─────────────────────────────────────────────────────

    _add_section_heading(document, "Proposed Team")

    _add_team_overview_table(document, entries)

    average_team_match = (
        round(
            sum(entry["match_score"] for entry in entries)
            / len(entries)
            * 100
        )
        if entries
        else 0
    )

    roles_with_gaps = sum(
        1 for entry in entries if entry["gap_count"] > 0
    )

    _add_team_metrics(
        document,
        average_team_match,
        roles_with_gaps,
        len(entries),
    )

    # ── Team AI assessment ────────────────────────────────────────────────

    _add_section_heading(document, "Team Assessment")

    assessment = document.add_paragraph(
        team_summary
        or "AI-generated team assessment was unavailable for this export."
    )

    assessment.paragraph_format.line_spacing = 1.15

    for run in assessment.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(DARK_GREY)

    document.add_page_break()

    # ── Individual assignment profiles ───────────────────────────────────

    heading = document.add_paragraph()

    heading_run = heading.add_run("INDIVIDUAL ASSIGNMENT PROFILES")
    heading_run.bold = True
    heading_run.font.name = "Arial"
    heading_run.font.size = Pt(16)
    heading_run.font.color.rgb = RGBColor.from_string(DARK_GREY)

    for index, entry in enumerate(entries):
        employee = entry["employee"]

        if index > 0:
            document.add_page_break()

        _add_section_heading(
            document,
            entry["role_title"],
        )

        employee_name = document.add_paragraph()

        employee_name_run = employee_name.add_run(
            employee.get("name", "")
        )
        employee_name_run.bold = True
        employee_name_run.font.name = "Arial"
        employee_name_run.font.size = Pt(16)
        employee_name_run.font.color.rgb = RGBColor.from_string(DARK_GREY)

        employee_details = document.add_paragraph()

        detail_values = [
            employee.get("title"),
            employee.get("role_level"),
            employee.get("business_unit"),
            employee.get("location"),
        ]

        detail_run = employee_details.add_run(
            " | ".join(
                str(value)
                for value in detail_values
                if value
            )
        )
        detail_run.font.name = "Arial"
        detail_run.font.size = Pt(9)
        detail_run.font.color.rgb = RGBColor.from_string(MID_GREY)

        years = employee.get("years_experience")
        if years is not None:
            _add_small_label(
                document,
                "Experience",
                f"{years} years",
            )

        _add_fit_summary(document, entry)

        _add_section_heading(document, "Profile")

        profile = document.add_paragraph(
            employee.get("summary", "")
            or "No employee summary recorded."
        )
        profile.paragraph_format.line_spacing = 1.15

        _add_section_heading(
            document,
            "Capability Alignment",
        )

        _add_capability_table(
            document,
            entry["fit_report"],
        )

        _add_section_heading(
            document,
            "Assignment Rationale",
        )

        rationale = document.add_paragraph(
            entry.get("rationale")
            or "AI-generated assignment rationale was unavailable for this export."
        )
        rationale.paragraph_format.line_spacing = 1.15

        _add_section_heading(
            document,
            "Relevant Experience",
        )

        _add_small_label(
            document,
            "Project experience",
            _clean_join(
                employee.get("project_experience", []) or []
            ),
        )

        _add_small_label(
            document,
            "Industry experience",
            _clean_join(
                employee.get("industry_experience", []) or []
            ),
        )

        _add_small_label(
            document,
            "Certifications",
            _clean_join(
                employee.get("certifications", []) or []
            ),
        )

    # ── Methodology note ──────────────────────────────────────────────────

    document.add_page_break()

    _add_section_heading(
        document,
        "About This Report",
    )

    methodology = document.add_paragraph(
        "Role-match results are generated by comparing project capability "
        "requirements against recorded employee skills using the capability "
        "matching model. AI-generated assessments interpret the deterministic "
        "matching and capability-fit results and are intended to support "
        "staffing decisions alongside professional judgement and other "
        "relevant workforce information."
    )

    methodology.paragraph_format.line_spacing = 1.15

    attribution = document.add_paragraph(
        "This service uses the ESCO classification of the European Commission."
    )

    attribution.paragraph_format.space_before = Pt(10)

    for run in attribution.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID_GREY)

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output